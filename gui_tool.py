import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
import requests
import re
from docx import Document
from docx.shared import Inches
import threading

class MeiwenCrawlerGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("美文网爬虫工具(https://www.ruiwen.com)：作者：小庄-Python办公")
        self.root.geometry("900x700")
        
        # 请求头和cookies
        self.headers = {
            "Accept": "application/json, text/javascript, */*; q=0.01",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8,en-GB;q=0.7,en-US;q=0.6",
            "Connection": "keep-alive",
            "Referer": "https://www.ruiwen.com/dldoc/index.html?0.6917949378500122222&pos=c&url=/jingdianmeiwen/1001432.html&jk=b&ps=99&view=copy",
            "Sec-Fetch-Dest": "empty",
            "Sec-Fetch-Mode": "cors",
            "Sec-Fetch-Site": "same-origin",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/139.0.0.0 Safari/537.36 Edg/139.0.0.0",
            "X-Requested-With": "XMLHttpRequest",
            "sec-ch-ua": '"Not;A=Brand";v="99", "Microsoft Edge";v="139", "Chromium";v="139"',
            "sec-ch-ua-mobile": "?0",
            "sec-ch-ua-platform": '"Windows"'
        }
        self.cookies = {
            "showNum": "1"
        }
        
        self.organized_content = {}
        self.setup_ui()
    
    def setup_ui(self):
        # 主框架
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 配置网格权重
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(2, weight=1)
        
        # 输入区域
        input_frame = ttk.LabelFrame(main_frame, text="输入区域", padding="10")
        input_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        input_frame.columnconfigure(1, weight=1)
        
        ttk.Label(input_frame, text="文章编号:").grid(row=0, column=0, sticky=tk.W, padx=(0, 10))
        self.article_id_var = tk.StringVar()
        self.article_entry = ttk.Entry(input_frame, textvariable=self.article_id_var, width=30)
        self.article_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(0, 10))
        
        self.fetch_btn = ttk.Button(input_frame, text="获取文章", command=self.fetch_article_threaded)
        self.fetch_btn.grid(row=0, column=2, padx=(0, 10))
        
        # 状态标签
        self.status_var = tk.StringVar(value="请输入文章编号")
        self.status_label = ttk.Label(input_frame, textvariable=self.status_var, foreground="blue")
        self.status_label.grid(row=1, column=0, columnspan=3, sticky=tk.W, pady=(10, 0))
        
        # 操作按钮区域
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        self.copy_btn = ttk.Button(button_frame, text="复制内容", command=self.copy_content, state="disabled")
        self.copy_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        self.export_txt_btn = ttk.Button(button_frame, text="导出为TXT", command=self.export_txt, state="disabled")
        self.export_txt_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        self.export_word_btn = ttk.Button(button_frame, text="导出为Word", command=self.export_word, state="disabled")
        self.export_word_btn.pack(side=tk.LEFT)
        
        # 预览区域
        preview_frame = ttk.LabelFrame(main_frame, text="文章预览", padding="10")
        preview_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S))
        preview_frame.columnconfigure(0, weight=1)
        preview_frame.rowconfigure(0, weight=1)
        
        self.preview_text = scrolledtext.ScrolledText(preview_frame, wrap=tk.WORD, state="disabled")
        self.preview_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 绑定回车键
        self.article_entry.bind('<Return>', lambda e: self.fetch_article_threaded())
    
    def fetch_article_threaded(self):
        """在新线程中获取文章，避免界面卡顿"""
        thread = threading.Thread(target=self.fetch_article)
        thread.daemon = True
        thread.start()
    
    def fetch_article(self):
        """获取文章内容"""
        article_id = self.article_id_var.get().strip()
        if not article_id:
            messagebox.showwarning("警告", "请输入文章编号")
            return
        
        self.status_var.set("正在获取文章...")
        self.fetch_btn.config(state="disabled")
        
        try:
            url = f"https://www.ruiwen.com/pic/dldoc/page1/{article_id}.html.txt"
            response = requests.get(url, headers=self.headers, cookies=self.cookies, timeout=10)
            
            if response.status_code != 200:
                raise Exception(f"HTTP错误: {response.status_code}")
            
            data = response.json()
            if 'body' not in data:
                raise Exception("响应数据格式错误")
            
            text = data['body']
            self.organized_content = self.parse_content(text)
            
            if not self.organized_content:
                raise Exception("未找到有效内容")
            
            self.display_content()
            self.enable_buttons()
            self.status_var.set(f"成功获取文章 {article_id}")
            
        except requests.exceptions.Timeout:
            self.status_var.set("请求超时，请重试")
            messagebox.showerror("错误", "请求超时，请检查网络连接")
        except requests.exceptions.RequestException as e:
            self.status_var.set("网络请求失败")
            messagebox.showerror("错误", f"网络请求失败: {str(e)}")
        except Exception as e:
            self.status_var.set("获取文章失败")
            messagebox.showerror("错误", f"获取文章失败: {str(e)}")
        finally:
            self.fetch_btn.config(state="normal")
    
    def parse_content(self, text):
        """解析文章内容"""
        # 使用正则表达式分割文本，同时保留标签信息
        pattern = r'(<h2>.*?</h2>|<p>.*?</p>)'
        elements = re.findall(pattern, text, re.DOTALL)
        
        # 组织内容：按h2标题分组p标签内容
        organized_content = {}
        current_title = "前言"  # 第一个h2标签之前的内容
        
        for element in elements:
            if element.startswith('<h2>'):
                # 提取h2标题文本
                title_match = re.search(r'<h2>(.*?)</h2>', element, re.DOTALL)
                if title_match:
                    current_title = re.sub(r'\s+', ' ', title_match.group(1).strip())
                    organized_content[current_title] = []
            elif element.startswith('<p>'):
                # 提取p标签文本
                p_match = re.search(r'<p>(.*?)</p>', element, re.DOTALL)
                if p_match:
                    p_text = re.sub(r'\s+', ' ', p_match.group(1).strip())
                    # 清理HTML标签
                    p_text = re.sub(r'<[^>]+>', '', p_text)
                    if current_title not in organized_content:
                        organized_content[current_title] = []
                    organized_content[current_title].append(p_text)
        
        return organized_content
    
    def display_content(self):
        """在预览区域显示内容"""
        self.preview_text.config(state="normal")
        self.preview_text.delete(1.0, tk.END)
        
        for title, paragraphs in self.organized_content.items():
            self.preview_text.insert(tk.END, f"【{title}】\n", "title")
            for i, para in enumerate(paragraphs, 1):
                self.preview_text.insert(tk.END, f"\n{para}\n")
            self.preview_text.insert(tk.END, "\n" + "="*50 + "\n\n")
        
        # 配置标题样式
        self.preview_text.tag_config("title", font=("微软雅黑", 12, "bold"), foreground="#2E86AB")
        
        self.preview_text.config(state="disabled")
    
    def enable_buttons(self):
        """启用操作按钮"""
        self.copy_btn.config(state="normal")
        self.export_txt_btn.config(state="normal")
        self.export_word_btn.config(state="normal")
    
    def copy_content(self):
        """复制内容到剪贴板"""
        if not self.organized_content:
            messagebox.showwarning("警告", "没有可复制的内容")
            return
        
        content = self.format_content_for_text()
        self.root.clipboard_clear()
        self.root.clipboard_append(content)
        messagebox.showinfo("成功", "内容已复制到剪贴板")
    
    def format_content_for_text(self):
        """格式化内容为纯文本"""
        content = []
        for title, paragraphs in self.organized_content.items():
            content.append(f"【{title}】")
            for para in paragraphs:
                content.append(f"\n{para}")
            content.append("\n" + "="*50 + "\n")
        return "\n".join(content)
    
    def export_txt(self):
        """导出为TXT文件"""
        if not self.organized_content:
            messagebox.showwarning("警告", "没有可导出的内容")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")],
            title="保存TXT文件"
        )
        
        if filename:
            try:
                content = self.format_content_for_text()
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(content)
                messagebox.showinfo("成功", f"文件已保存到: {filename}")
            except Exception as e:
                messagebox.showerror("错误", f"保存文件失败: {str(e)}")
    
    def export_word(self):
        """导出为Word文件"""
        if not self.organized_content:
            messagebox.showwarning("警告", "没有可导出的内容")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")],
            title="保存Word文件"
        )
        
        if filename:
            try:
                doc = Document()
                
                for title, paragraphs in self.organized_content.items():
                    # 添加标题
                    heading = doc.add_heading(title, level=1)
                    heading.alignment = 1  # 居中对齐
                    
                    # 添加段落
                    for para in paragraphs:
                        p = doc.add_paragraph(para)
                        p.alignment = 0  # 左对齐
                    
                    # 添加分隔线
                    doc.add_paragraph("="*50)
                
                doc.save(filename)
                messagebox.showinfo("成功", f"Word文档已保存到: {filename}")
            except Exception as e:
                messagebox.showerror("错误", f"保存Word文档失败: {str(e)}")

def main():
    root = tk.Tk()
    app = MeiwenCrawlerGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()